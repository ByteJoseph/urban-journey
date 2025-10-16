import streamlit as st
from docx import Document
from io import BytesIO
import os

st.set_page_config(page_title="Word Template Filler", layout="centered")

st.title("📝 Word Template Auto-Filler (Local File)")
st.caption("Automatically fill placeholders in your local Word template (no upload needed).")

# === Path to your local Word file ===
template_path = "tg.docx"  # ensure this file exists in the same directory as this script

if not os.path.exists(template_path):
    st.error(f"❌ '{template_path}' not found in the current directory.")
else:
    # Input fields for placeholders
    st.subheader("Fill in the details")

    name = st.text_input("Name", "JOSEPH JOSEPH")
    date = st.text_input("Date (dd/mm/yyyy)", "16/10/2025")
    title = st.text_input("Title", "POLYNOMIAL ADDITION USING ARRAY")
    roll = st.text_input("Roll No", "23")
    exp_no = st.text_input("Experiment No", "13")
    program = st.text_area("Program Code", """#include <stdio.h>

int main() {
    int n1, n2;
    printf("Enter number of terms in first polynomial: ");
    scanf("%d", &n1);
    printf("Enter number of terms in second polynomial: ");
    scanf("%d", &n2);
    printf("Polynomial addition successful!");
    return 0;
}""", height=220)

    output = st.text_area("Program Output", """Enter number of terms in first polynomial: 2
Enter coefficient of term 1: 3
Enter exponent of term 1: 2
Enter coefficient of term 2: 4
Enter exponent of term 2: 0
Enter number of terms in second polynomial: 2
Enter coefficient of term 1: 5
Enter exponent of term 1: 1
Enter coefficient of term 2: 6
Enter exponent of term 2: 0
First Polynomial: 3x^2 + 4x^0
Second Polynomial: 5x^1 + 6x^0
Result : 3x^2 + 5x^1 + 10x^0""", height=200)

    if st.button("🔧 Generate Word File"):
        # Load the local Word document
        doc = Document(template_path)

        # Placeholder replacements
        replacements = {
            "{{NAME}}": name,
            "{{dd/mm/yyyy}}": date,
            "{{TITLE}}": title,
            "{{roll}}": roll,
            "{{no}}": exp_no,
            "{{program}}": program,
            "{{output}}": output,
        }

        # Replace text in paragraphs
        for p in doc.paragraphs:
            for key, value in replacements.items():
                if key in p.text:
                    p.text = p.text.replace(key, value)

        # Replace inside tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in p.text:
                                p.text = p.text.replace(key, value)

        # Save edited document to memory
        output_stream = BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)

        st.success("✅ Document generated successfully!")

        st.download_button(
            label="📥 Download Edited Word File",
            data=output_stream,
            file_name="Edited_Template.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
